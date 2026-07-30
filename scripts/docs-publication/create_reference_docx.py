from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

PURPLE = RGBColor(109, 74, 176)
DARK = RGBColor(23, 21, 31)
MUTED = RGBColor(98, 93, 111)
LIGHT = "F0EBF8"
PAGE_WIDTH_DXA = 9026


def set_font(run, name: str, size: float, color: RGBColor = DARK, bold: bool = False) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instruction, separate, text, end):
        run._r.append(element)


def configure_style(document: Document, name: str, font_size: float, *,
                    color: RGBColor = DARK, bold: bool = False,
                    before: float = 0, after: float = 6,
                    line_spacing: float = 1.25) -> None:
    style = document.styles[name]
    set_font(style.element.get_or_add_rPr(), "Calibri", font_size, color, bold)
    style.font.name = "Calibri"
    style.font.size = Pt(font_size)
    style.font.color.rgb = color
    style.font.bold = bold
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = line_spacing


def configure_run_properties(style, font_size: float, color: RGBColor, bold: bool) -> None:
    style.font.name = "Calibri"
    style.font.size = Pt(font_size)
    style.font.color.rgb = color
    style.font.bold = bold
    style.element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    style.element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")


def build_reference(output_path: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.right_margin = Cm(1.8)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(1.8)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    normal = document.styles["Normal"]
    configure_run_properties(normal, 10.5, DARK, False)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = document.styles["Title"]
    configure_run_properties(title, 30, DARK, True)
    title.paragraph_format.space_before = Pt(70)
    title.paragraph_format.space_after = Pt(10)

    subtitle = document.styles["Subtitle"]
    configure_run_properties(subtitle, 15, MUTED, False)
    subtitle.paragraph_format.space_after = Pt(24)

    heading_tokens = {
        "Heading 1": (20, PURPLE, 18, 10),
        "Heading 2": (15, PURPLE, 14, 7),
        "Heading 3": (12, RGBColor(52, 43, 69), 10, 5),
        "Heading 4": (10.5, DARK, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = document.styles[name]
        configure_run_properties(style, size, color, True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("Caption", "Image Caption"):
        if name not in document.styles:
            document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style = document.styles[name]
        configure_run_properties(style, 8.75, MUTED, False)
        style.font.italic = True
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style.paragraph_format.space_before = Pt(4)
        style.paragraph_format.space_after = Pt(8)

    for name in ("Source Code", "Verbatim Char"):
        if name not in document.styles:
            style_type = WD_STYLE_TYPE.CHARACTER if name == "Verbatim Char" else WD_STYLE_TYPE.PARAGRAPH
            document.styles.add_style(name, style_type)
        style = document.styles[name]
        configure_run_properties(style, 8.5, DARK, False)
        style.font.name = "Cascadia Mono"
        style.element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Cascadia Mono")
        style.element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Cascadia Mono")
        if style.type == WD_STYLE_TYPE.PARAGRAPH:
            style.paragraph_format.space_before = Pt(4)
            style.paragraph_format.space_after = Pt(7)
            style.paragraph_format.keep_together = True

    if "Callout" not in document.styles:
        document.styles.add_style("Callout", WD_STYLE_TYPE.PARAGRAPH)
    callout = document.styles["Callout"]
    configure_run_properties(callout, 10, DARK, False)
    callout.paragraph_format.left_indent = Cm(0.5)
    callout.paragraph_format.right_indent = Cm(0.5)
    callout.paragraph_format.space_before = Pt(6)
    callout.paragraph_format.space_after = Pt(8)

    if "Hyperlink" not in document.styles:
        document.styles.add_style("Hyperlink", WD_STYLE_TYPE.CHARACTER)
    hyperlink = document.styles["Hyperlink"]
    configure_run_properties(hyperlink, 10.5, PURPLE, False)
    hyperlink.font.underline = True

    header = section.header.paragraphs[0]
    header.text = "TRACKLY  |  TECHNICAL DOCUMENTATION"
    set_font(header.runs[0], "Calibri", 8.5, MUTED, True)
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT

    footer = section.footer.paragraphs[0]
    footer.add_run("Trackly  |  Version 1.0                                      ")
    set_font(footer.runs[0], "Calibri", 8, MUTED)
    add_page_field(footer)

    document.add_paragraph("Trackly Technical Documentation", style="Title")
    document.add_paragraph(
        "Architecture, Development, Features, and Operations",
        style="Subtitle",
    )
    metadata = document.add_paragraph()
    metadata.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = metadata.add_run("Version 1.0\nFahmy Akhmad Firdaus\n2026")
    set_font(run, "Calibri", 11, MUTED)
    metadata.add_run().add_break(WD_BREAK.PAGE)

    document.add_heading("Reference Heading 1", level=1)
    document.add_heading("Reference Heading 2", level=2)
    document.add_heading("Reference Heading 3", level=3)
    document.add_heading("Reference Heading 4", level=4)
    document.add_paragraph("Reference body paragraph for publication styling.")
    document.add_paragraph("const status: string = 'ready';", style="Source Code")
    document.add_paragraph("Figure 1.1 - Reference caption", style="Caption")
    document.add_paragraph("Important implementation note.", style="Callout")

    table = document.add_table(rows=2, cols=2)
    table.autofit = False
    widths = (Cm(4), Cm(11.92))
    for row_index, row in enumerate(table.rows):
        for cell_index, cell in enumerate(row.cells):
            cell.width = widths[cell_index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = ("Field", "Value")[cell_index] if row_index == 0 else ("Status", "Ready")[cell_index]
            if row_index == 0:
                set_cell_shading(cell, LIGHT)
                for run in cell.paragraphs[0].runs:
                    run.bold = True

    document.add_section(WD_SECTION.NEW_PAGE)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


if __name__ == "__main__":
    if len(sys.argv) != 2:
        raise SystemExit("Usage: create_reference_docx.py <output.docx>")
    build_reference(Path(sys.argv[1]).resolve())
